#không dùng tới
from docx import Document
from docx.shared import Pt,Inches,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.shared import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.enum.text import WD_BREAK
import copy

# hàm tạo 1 table với số lượng cột cho trước
def create_table(doc, num_rows, num_cols, border_val='single', border_sz='0.25'):
    """
    Tạo một bảng mới với số hàng và cột được chỉ định, có thể thiết lập border.

    :param doc: Đối tượng Document từ python-docx
    :param num_rows: Số lượng hàng trong bảng
    :param num_cols: Số lượng cột trong bảng
    :param border_val: Kiểu border (mặc định 'single')
    :param border_sz: Độ dày border (mặc định '5')
    :return: Đối tượng Table mới được tạo
    """
    table = doc.add_table(rows=num_rows, cols=num_cols)
    set_table_borders(table, border_val, border_sz)
    return table

# Hàm xử lý RPno trong header
def replace_in_header_table(doc, row, col, key, value, section_index=0, table_index=0):
    header = doc.sections[section_index].header
    table = header.tables[table_index]
    cell = table.cell(row, col)
    cell.text = cell.text.replace(key, value)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Arial'
    return doc


# Hàm đưa dữ liệu vào file word
def replace_text_in_word(doc,dic):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dic.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial MT'  
                                run.font.size = Pt(9) 

# tìm table index với tiêu đề
def find_tables_with_title(doc,table_title):
    """Nếu tìm không thấy trả về -1"""
    found = -1
    for index, table in enumerate(doc.tables):
        tbl_element = table._element
        tblPr = tbl_element.xpath('./w:tblPr')
        if tblPr:
            tblPr = tblPr[0]
            tblCaption = tblPr.xpath('./w:tblCaption')
            if tblCaption:
                alt_text = tblCaption[0].get(qn('w:val'))
                if alt_text == table_title:
                    found = index
                    break
    return found

# thêm dòng tùy ý cho doc table
def add_rows_to_table(doc, table_index, rows_to_add):
    if 0 <= table_index < len(doc.tables):
        table = doc.tables[table_index]
        for _ in range(rows_to_add):
            table.add_row()
    return doc

# merge các dòng, cột tùy ý cho doc table
def merge_table_cells(doc, table_index, start_row, end_row, start_col, end_col):
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError("Invalid table index")
    table = doc.tables[table_index]
    
    if start_row < 0 or end_row >= len(table.rows) or start_col < 0 or end_col >= len(table.columns):
        raise ValueError("Invalid cell range")
    
    table.cell(start_row, start_col).merge(table.cell(end_row, end_col))
    
    return doc

# điền dữ liệu vào từng ô trong table với định dạng
def fill_table_cell(doc, table_index, row, col, text_parts, 
                    alignment=None, vertical_alignment=None,
                    border_color=None, border_size=None, border_style=None, border_sides = ['top', 'left', 'bottom', 'right']):
    """
    Fill a specific cell in a table with formatted text and cell properties.
    
    :param doc: Document object
    :param table_index: Index of the table (0-based)
    :param row: Row index of the cell (0-based)
    :param col: Column index of the cell (0-based)
    :param text_parts: List of dicts, each containing text and its formatting
    :param alignment: String, text alignment ('LEFT', 'CENTER', 'RIGHT', 'JUSTIFY')
    :param vertical_alignment: String, vertical_alignment ('TOP', 'CENTER', 'BOTTOM')
    :param border_color: Tuple of (R, G, B) values for border color
    :param border_size: Integer, border width in points
    :param border_style: String, border style ('single', 'dash', 'dotted', etc.)
    
        Giá trị	Mô tả boder
    'single'	Đường liền đơn (mặc định)
    'thick'	Đường liền dày
    'double'	Đường đôi
    'dotted'	Đường chấm
    'dashed'	Đường đứt
    'dotDash'	Đường chấm-gạch
    'dotDotDash'	Đường chấm-chấm-gạch
    'triple'	Ba đường liền
    'thinThickSmallGap'	Đường mỏng-dày cách nhỏ
    'thickThinSmallGap'	Đường dày-mỏng cách nhỏ
    'thinThickThinSmallGap'	Đường mỏng-dày-mỏng cách nhỏ
    'nil'	Không có viền
    
    Example usage:
    fill_table_cell(doc, 0, 1, 2, [
        {'text': 'Hello ', 'bold': True, 'font_name': 'Arial', 'font_size': 12, 'color': (255, 0, 0)},
        {'text': 'World!', 'italic': True, 'font_name': 'Calibri', 'font_size': 14, 'color': (0, 255, 0)},
        {'text': '\nNew line', 'underline': True, 'font_name': 'Times New Roman', 'font_size': 10, 'highlight': 'YELLOW'}
    ], alignment='CENTER')
    """
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError("Invalid table index")
    
    table = doc.tables[table_index]
    
    if row < 0 or row >= len(table.rows) or col < 0 or col >= len(table.columns):
        raise ValueError("Invalid cell coordinates")
    
    cell = table.cell(row, col)
    cell.text = ""  # Clear existing content
    
    paragraph = cell.paragraphs[0]
    # Set default indents to 0
    paragraph.paragraph_format.first_line_indent = 0
    paragraph.paragraph_format.left_indent = 0
    paragraph.paragraph_format.hanging_indent = 0
    
    for part in text_parts:
        part['text'] = str(part['text'])
        if '\n' in part['text']:
            # Split the text into lines and create new paragraphs for each line
            lines = part['text'].split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    paragraph = cell.add_paragraph()
                    # Set default indents to 0 for new paragraph
                    paragraph.paragraph_format.first_line_indent = 0
                    paragraph.paragraph_format.left_indent = 0
                    paragraph.paragraph_format.hanging_indent = 0
                run = paragraph.add_run(line)
                apply_run_formatting(run, part)
                # Apply alignment to each paragraph
                if alignment:
                    try:
                        paragraph.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment.upper())
                    except AttributeError:
                        print(f"Warning: Alignment '{alignment}' not recognized. Skipping alignment.")
        else:
            run = paragraph.add_run(part['text'])
            apply_run_formatting(run, part)
    
    # Apply alignment 
    if alignment and not paragraph.alignment:
        try:
            paragraph.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment.upper())
        except AttributeError:
            print(f"Warning: Alignment '{alignment}' not recognized. Skipping alignment.")
    
    # Cell formatting
    if vertical_alignment:
        try:
            cell.vertical_alignment = getattr(WD_CELL_VERTICAL_ALIGNMENT, vertical_alignment.upper())
        except AttributeError:
            print(f"Warning: vertical_alignment '{vertical_alignment}' not recognized. Skipping vertical_alignment.")
    
   # Border formatting 
    if border_color or border_size or border_style:
        tcPr = cell._element.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell._element.append(tcPr)
        
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
        
        border_style_map = {
            'single': 'single',
            'thick': 'thick',
            'double': 'double',
            'dotted': 'dotted',
            'dashed': 'dashed',
            'dotDash': 'dotDash',
            'dotDotDash': 'dotDotDash',
            'triple': 'triple',
            'thinThickSmallGap': 'thinThickSmallGap',
            'thickThinSmallGap': 'thickThinSmallGap',
            'thinThickThinSmallGap': 'thinThickThinSmallGap',
            'nil': 'nil'
        }
        
        style_val = border_style_map.get(border_style, 'single')
        
        for border_side in border_sides:
            border = OxmlElement(f'w:{border_side}')
            borders.append(border)
            
            border.set(qn('w:val'), style_val)
            
            if border_size:
                border.set(qn('w:sz'), str(border_size * 8))
                
            if border_color:
                r, g, b = border_color
                color_str = f'{r:02x}{g:02x}{b:02x}'
                border.set(qn('w:color'), color_str)
    return doc

# format 1 range trong table word
def apply_format_to_table_range(doc, table_index, start_row, end_row, start_col, end_col, format_dict):
    """
    Áp dụng định dạng cho một phạm vi cụ thể trong bảng của tài liệu.

    :param doc: Đối tượng Document từ python-docx
    :param table_index: Chỉ số của bảng trong tài liệu (0-based)
    :param start_row: Chỉ số hàng bắt đầu (0-based)
    :param end_row: Chỉ số hàng kết thúc (0-based, bao gồm)
    :param start_col: Chỉ số cột bắt đầu (0-based)
    :param end_col: Chỉ số cột kết thúc (0-based, bao gồm)
    :param format_dict: Dictionary chứa thông tin định dạng

    Ví dụ sử dụng format_dict:
        format_dict = {
            'bold': True,
            'italic': False,
            'underline': True,
            'font_name': 'Arial',
            'font_size': 12,
            'color': (255, 0, 0),         # Đỏ
            'highlight': 'YELLOW',        # Tô nền chữ vàng
            'SHADING': (220, 230, 241)    # Màu nền ô (RGB)
        }
    """
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"Table index {table_index} is out of range. Document has {len(doc.tables)} tables.")

    table = doc.tables[table_index]

    if start_row < 0 or end_row >= len(table.rows) or start_col < 0 or end_col >= len(table.columns):
        raise ValueError("Row or column indices are out of range for the specified table.")

    for row in range(start_row, end_row + 1):
        for col in range(start_col, end_col + 1):
            cell = table.cell(row, col)
            
            # Áp dụng shading nếu được chỉ định
            if 'SHADING' in format_dict:
                shading_color = format_dict['SHADING']
                set_cell_shading(cell, shading_color)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    apply_run_formatting(run, format_dict)

# định dạng text
def apply_run_formatting(run, format_dict):
    run.bold = format_dict.get('bold', False)
    run.italic = format_dict.get('italic', False)
    run.underline = format_dict.get('underline', False)
    if 'font_name' in format_dict:
        run.font.name = format_dict['font_name']
    if 'font_size' in format_dict:
        run.font.size = Pt(format_dict['font_size'])
    if 'color' in format_dict:
        run.font.color.rgb = RGBColor(*format_dict['color'])
    if 'highlight' in format_dict and format_dict['highlight'] is not None:
        try:
            run.font.highlight_color = getattr(WD_COLOR_INDEX, format_dict['highlight'].upper())
        except AttributeError:
            print(f"Warning: Highlight color '{format_dict['highlight']}' not recognized. Skipping highlight.")
    if 'superscript' in format_dict:
        run.font.superscript = format_dict['superscript']
# định dạng shading 
def set_cell_shading(cell, color):
    """
    Set background color for a table cell.

    :param cell: Cell object from python-docx
    :param color: Tuple (R, G, B) for background color
    """
    hex_color = rgb_to_hex(color)
    shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def rgb_to_hex(rgb):
    """Convert an RGB tuple to a hex string without '#'."""
    return '%02X%02X%02X' % rgb

# xóa table theo table index và có thể xóa thêm các dòng trống sau bảng
def delete_table(doc, table_index, extra_lines=0):
    """
    Xóa bảng theo chỉ số và xóa thêm extra_lines đoạn văn sau bảng (nếu có).
    """
    table = doc.tables[table_index]
    tbl_elem = table._element
    parent = tbl_elem.getparent()
    idx = parent.index(tbl_elem)
    parent.remove(tbl_elem)
    
    # Xóa thêm các dòng trống (w:p) sau vị trí bảng vừa xóa
    for _ in range(extra_lines):
        if idx < len(parent):
            next_elem = parent[idx]
            if next_elem.tag.endswith('}p'):
                parent.remove(next_elem)
            else:
                break
    return doc

# lấy ra danh sách chiều rộng của 1 hàng theo table index
def get_row_cell_widths(doc, table_index, row_index):
    """
    Trả về danh sách độ rộng của các ô (tính bằng twips) cho một hàng cụ thể trong bảng.
    """
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError("Chỉ số bảng không hợp lệ")
    table = doc.tables[table_index]
    if row_index < 0 or row_index >= len(table.rows):
        raise ValueError("Chỉ số hàng không hợp lệ")
    widths = []
    for cell in table.rows[row_index].cells:
        tc = cell._tc
        tcPr = tc.tcPr
        width_elem = tcPr.find(qn('w:tcW'))
        if width_elem is not None and width_elem.get(qn('w:w')) is not None:
            width = int(width_elem.get(qn('w:w')))
            widths.append(width)
        else:
            widths.append(None)
    return widths

# chèn số lượng dòng trống sau table chỉ định
def insert_blank_paragraphs_after_table(doc, table_index, num_paragraphs=1):
    """
    Chèn các đoạn văn trống sau bảng được chỉ định.
    """
    table = doc.tables[table_index] 
    tbl = table._element
    parent = tbl.getparent()
    idx = parent.index(tbl)
    for _ in range(num_paragraphs):
        p = OxmlElement('w:p')
        parent.insert(idx + 1, p)
        idx += 1
    return idx, parent

# tạo 1 bảng với danh sách độ rộng cho trướng
def create_table_with_style_and_widths(doc, style, cell_widths):
    """
    Tạo một bảng mới với kiểu và độ rộng ô được chỉ định.
    """
    cols = len(cell_widths)
    table = doc.add_table(rows=1, cols=cols)
    table.style = style
    tbl = table._element
    tblPr = tbl.xpath('./w:tblPr')[0]
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    total_width = sum([w for w in cell_widths if w is not None])
    tblW.set(qn('w:w'), str(total_width))
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')
    for i, width in enumerate(cell_widths):
        if width is not None:
            emu_width = int(width / 1440 * 914400)
            table.columns[i].width = emu_width
            for cell in table.columns[i].cells:
                cell.width = emu_width
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(width))
                tcW.set(qn('w:type'), 'dxa')
    return table

# thiết lập đường viền cho border table
def set_table_borders(table, border_val, border_sz):
    """
    Thiết lập đường viền cho bảng.
    """
    tbl_elem = table._element
    tblPr = tbl_elem.xpath('./w:tblPr')[0]
    borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), border_val)
        border.set(qn('w:sz'), border_sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)

# chèn table vào vị trí cụ thể
def insert_table_after(parent, idx, table):
    """
    Chèn bảng vào vị trí cụ thể trong tài liệu.
    """
    parent.remove(table._element)
    parent.insert(idx + 1, table._element)

# thiết lập độ rộng table 100%
def set_table_width_100_percent(doc, table_index):
    """
    Đặt độ rộng của bảng thành 100% của trang.
    """
    table = doc.tables[table_index]
    tbl_pr = table._element.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)

    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), '5000')  # 5000 = 100%
    return doc

# ------------------..............--------------------------
# Thêm 1 bảng mới vào sau 1 bảng cụ thể
def add_table_after_table_index(doc, table_index, cell_widths, num_paragraphs=1, border_val='single', border_sz='5', table_width = "100%"):
    """
    Thêm một bảng mới sau bảng được chỉ định bởi chỉ số.
    """
    cell_widths = [int(width*1440) for width in cell_widths] 
    src_table = doc.tables[table_index]
    idx, parent = insert_blank_paragraphs_after_table(doc, table_index, num_paragraphs)
    new_table = create_table_with_style_and_widths(doc, src_table.style, cell_widths)
    set_table_borders(new_table, border_val, border_sz)
    insert_table_after(parent, idx, new_table)
    if table_width == "100%":
        set_table_width_100_percent(doc, table_index+1)

# ------------------..............--------------------------
# hàm chèn sang trang từ table
def insert_page_break_after_table(doc, table_index):
    table = doc.tables[table_index]
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

# hàm chèn nội dung từ file docx vào file docx
def insert_content_from_docx(template_path, doc, insert_after_text=None, insert_at_start=False, 
                           insert_after_table_index=None, source_section_index=None):
    """
    Chèn nội dung từ file Word template_path vào file Word doc.
    
    Tham số vị trí chèn:
    - insert_at_start=True: Chèn vào đầu file đích.
    - insert_after_text='nội dung': Chèn sau đoạn văn bản chứa 'nội dung' trong file đích.
    - insert_after_table_index=1: Chèn sau bảng thứ 2 (bắt đầu từ 0) trong file đích.
    - Nếu không truyền gì, mặc định chèn vào cuối file đích.
    
    Tham số nguồn dữ liệu:
    - source_section_index=0: Chỉ lấy nội dung từ section đầu tiên (bắt đầu từ 0).
    - Nếu không truyền, lấy toàn bộ nội dung từ template.
    """
    src_doc = Document(template_path)
    dst_doc = doc
    dst_body = dst_doc._element.body

    if source_section_index is not None:
        if source_section_index < len(src_doc.sections):
            section = src_doc.sections[source_section_index]
            content_elements = get_section_content(src_doc, source_section_index)
        else:
            print(f"Warning: Section index {source_section_index} không tồn tại. Sử dụng toàn bộ document.")
            content_elements = list(src_doc._element.body)
    else:
        content_elements = list(src_doc._element.body)

    elements_to_insert = [copy.deepcopy(el) for el in content_elements 
                     if not el.tag.endswith('sectPr')]
    insert_pos = determine_insert_position(dst_body, insert_at_start, insert_after_text, insert_after_table_index)

    for element in elements_to_insert:
        dst_body.insert(insert_pos, element)
        insert_pos += 1

# hàm lấy nội dung của một section cụ thể từ document
def get_section_content(doc, section_index):
    """
    Lấy nội dung của một section cụ thể từ document.
    """
    if section_index >= len(doc.sections):
        return []
    
    body_elements = list(doc._element.body)
    section_breaks = []
    
    for i, element in enumerate(body_elements):
        if element.tag.endswith('sectPr') or has_section_break(element):
            section_breaks.append(i)
    
    if not section_breaks:
        return body_elements
    
    if section_index == 0:
        start_idx = 0
        end_idx = section_breaks[0] if section_breaks else len(body_elements)
    elif section_index < len(section_breaks):
        start_idx = section_breaks[section_index - 1] + 1
        end_idx = section_breaks[section_index]
    else:
        start_idx = section_breaks[-1] + 1
        end_idx = len(body_elements)
    
    return body_elements[start_idx:end_idx]

# hàm kiểm tra xem element có chứa section break không
def has_section_break(element):
    """
    Kiểm tra xem element có chứa section break không.
    """
    for child in element.iter():
        if child.tag.endswith('br') and child.get('type') == 'page':
            return True
        if child.tag.endswith('sectPr'):
            return True
    return False

# hàm xác định vị trí chèn trong document đích
def determine_insert_position(dst_body, insert_at_start, insert_after_text, insert_after_table_index):
    """
    Xác định vị trí chèn trong document đích.
    """
    if insert_at_start:
        return 0
    elif insert_after_table_index is not None:
        tbl_count = -1
        for idx, el in enumerate(dst_body):
            if el.tag.endswith('tbl'):
                tbl_count += 1
                if tbl_count == insert_after_table_index:
                    return idx + 1
        return len(dst_body)
    elif insert_after_text:
        for idx, el in enumerate(dst_body):
            if el.tag.endswith('p'):
                texts = [node.text for node in el.iter() if node.text]
                para_text = ''.join(texts)
                if para_text and insert_after_text in para_text:
                    return idx + 1
        return len(dst_body)
    else:
        return len(dst_body)

# hàm hiển thị thông tin về các section trong document
def list_sections_info(docx_path):
    """
    Hiển thị thông tin về các section trong document.
    """
    doc = Document(docx_path)
    print(f"Document có {len(doc.sections)} section(s):")
    
    for i, section in enumerate(doc.sections):
        print(f"Section {i}:")
        print(f"  - Orientation: {'Portrait' if section.orientation == 0 else 'Landscape'}")
        print(f"  - Page width: {section.page_width}")
        print(f"  - Page height: {section.page_height}")
        
        content = get_section_content(doc, i)
        para_count = sum(1 for el in content if el.tag.endswith('p'))
        table_count = sum(1 for el in content if el.tag.endswith('tbl'))
        print(f"  - Paragraphs: {para_count}")
        print(f"  - Tables: {table_count}")
        print()

# --------------------------------------------------------------
# Hàm chèn table từ file Word khác vào document hiện tại
def insert_table_from_docx(doc, table_index,template_path = None,  insert_after_text=None, 
                          insert_at_start=False, insert_after_table_index=None, 
                          source_section_index=None, insert_num_line=1):
    """
    Chèn một table cụ thể từ file Word template_path vào file Word doc.
    
    Tham số nguồn dữ liệu:
    - table_index: Index của table cần lấy từ template (bắt đầu từ 0).
    - source_section_index: Index của section chứa table (nếu không truyền thì tìm trong toàn bộ document).
    
    Tham số vị trí chèn:
    - insert_at_start=True: Chèn vào đầu file đích.
    - insert_after_text='nội dung': Chèn sau đoạn văn bản chứa 'nội dung' trong file đích.
    - insert_after_table_index=1: Chèn sau bảng thứ 2 (bắt đầu từ 0) trong file đích.
    - insert_num_line: Số dòng trống cần thêm trước table (mặc định 1).
    - Nếu không truyền gì, mặc định chèn vào cuối file đích.
    
    Returns:
    - True nếu chèn thành công, False nếu không tìm thấy table.
    """
    if template_path is None:
        src_doc = doc
    else:
        src_doc = Document(template_path)
    dst_doc = doc
    dst_body = dst_doc._element.body

    if source_section_index is not None:
        if source_section_index < len(src_doc.sections):
            content_elements = get_section_content(src_doc, source_section_index)
        else:
            print(f"Warning: Section index {source_section_index} không tồn tại. Tìm trong toàn bộ document.")
            content_elements = list(src_doc._element.body)
    else:
        content_elements = list(src_doc._element.body)

    table_element = find_table_by_index(content_elements, table_index)
    
    if table_element is None:
        return False

    table_copy = copy.deepcopy(table_element)
    insert_pos = determine_insert_position(dst_body, insert_at_start, insert_after_text, insert_after_table_index)
    
    # Thêm các paragraph trống trước table
    for i in range(insert_num_line):
        empty_para = create_empty_paragraph()
        dst_body.insert(insert_pos + i, empty_para)
    
    # Chèn table sau các paragraph trống
    dst_body.insert(insert_pos + insert_num_line, table_copy)
    
    return True

def create_empty_paragraph():
    """
    Tạo một paragraph trống ở mức XML element.
    """
    para = OxmlElement('w:p')
    para.append(OxmlElement('w:r'))
    return para

# hàm tìm table theo index
def find_table_by_index(elements, table_index):
    """
    Tìm table theo index trong danh sách elements.
    """
    table_count = -1
    for element in elements:
        if element.tag.endswith('tbl'):
            table_count += 1
            if table_count == table_index:
                return element
    return None

# Hàm hiển thị thông tin về các table trong document hoặc section cụ thể
def list_tables_info(docx_path, section_index=None):
    """
    Hiển thị thông tin về các table trong document hoặc section cụ thể.
    """
    doc = Document(docx_path)
    
    if section_index is not None:
        if section_index >= len(doc.sections):
            print(f"Error: Section index {section_index} không tồn tại.")
            return
        print(f"Tables trong Section {section_index}:")
        content_elements = get_section_content(doc, section_index)
    else:
        print("Tables trong toàn bộ document:")
        content_elements = list(doc._element.body)
    
    table_count = -1
    for i, element in enumerate(content_elements):
        if element.tag.endswith('tbl'):
            table_count += 1
            # Lấy thông tin chi tiết về table
            table_info = get_table_info(doc, element)
            print(f"  Table {table_count}:")
            print(f"    - Vị trí trong document: element {i}")
            print(f"    - Số hàng: {table_info['rows']}")
            print(f"    - Số cột: {table_info['cols']}")
            print(f"    - Nội dung hàng đầu: {table_info['first_row']}")
            print()
    
    if table_count == -1:
        print("  Không có table nào.")

# Hàm lấy thông tin chi tiết về một table element
def get_table_info(doc, table_element):
    """
    Lấy thông tin chi tiết về một table element.
    """
    # Đếm số hàng và cột
    rows = table_element.xpath('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    cols = 0
    first_row_text = ""
    
    if rows:
        first_row = rows[0]
        cells = first_row.xpath('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        cols = len(cells)
        
        # Lấy text từ các cell của hàng đầu
        cell_texts = []
        for cell in cells[:3]:  # Chỉ lấy 3 cell đầu để preview
            texts = [node.text for node in cell.iter() if node.text and node.text.strip()]
            cell_text = ' '.join(texts).strip()
            if cell_text:
                cell_texts.append(cell_text[:20] + "..." if len(cell_text) > 20 else cell_text)
        first_row_text = " | ".join(cell_texts)
    
    return {
        'rows': len(rows),
        'cols': cols,
        'first_row': first_row_text or "Không có nội dung"
    }

# Hàm tiện ích để chèn nhiều table cùng lúc
def insert_multiple_tables_from_docx(template_path, doc, table_indices, insert_after_text=None, 
                                   insert_at_start=False, insert_after_table_index=None, 
                                   source_section_index=None, add_spacing=True):
    """
    Chèn nhiều table cùng lúc từ template.
    
    Tham số:
    - table_indices: List các index của table cần chèn [0, 2, 3].
    - add_spacing: Thêm paragraph trống giữa các table (mặc định True).
    """
    success_count = 0
    
    for i, table_idx in enumerate(table_indices):
        # Tính toán vị trí chèn cho table tiếp theo
        if i == 0:
            # Table đầu tiên dùng vị trí được chỉ định
            current_insert_after_table = insert_after_table_index
            current_insert_after_text = insert_after_text
            current_insert_at_start = insert_at_start
        else:
            # Các table tiếp theo chèn sau table vừa chèn
            current_insert_after_table = None
            current_insert_after_text = None
            current_insert_at_start = False
            # Tính toán lại vị trí dựa trên số table đã chèn
        
        success = insert_table_from_docx(
            template_path, doc, table_idx,
            insert_after_text=current_insert_after_text,
            insert_at_start=current_insert_at_start,
            insert_after_table_index=current_insert_after_table,
            source_section_index=source_section_index
        )
        
        if success:
            success_count += 1
            
            # Thêm paragraph trống giữa các table nếu cần
            if add_spacing and i < len(table_indices) - 1:
                add_empty_paragraph(doc)
    
    print(f"Đã chèn thành công {success_count}/{len(table_indices)} table(s)")
    return success_count

def add_empty_paragraph(doc):
    """
    Thêm một paragraph trống vào cuối document.
    """
    doc.add_paragraph("")
    
# --------------------------Hàm chèn 1 paragraph text vào sau table--------------------------
def insert_paragraph_after_table(doc, table_index, text, font_name='Arial', font_size=8, bold=False):
    """
    Chèn một paragraph text vào sau table chỉ định.
    """
    table = doc.tables[table_index]
    new_paragraph = doc.add_paragraph(text)
    run = new_paragraph.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if bold:
        run.bold = bold
    table._element.addnext(new_paragraph._element)
    
    # --------------------------đọc dữ liệu header table từ file docx từ cell cụ thể--------------------------
def read_header_table_cell(doc, section_index=0, table_index=0, row=0, col=0):
    """
    Đọc dữ liệu từ một ô cụ thể trong header table của section chỉ định.
    
    :param doc: Đối tượng Document từ python-docx
    :param section_index: Chỉ số của section (mặc định 0)
    :param table_index: Chỉ số của bảng trong header (mặc định 0)
    :param row: Chỉ số hàng của ô cần đọc (mặc định 0)
    :param col: Chỉ số cột của ô cần đọc (mặc định 0)
    :return: Nội dung văn bản trong ô
    """
    header = doc.sections[section_index].header
    table = header.tables[table_index]
    cell = table.cell(row, col)
    return cell.text.strip()

# --------------------------lấy danh sách index của table theo alt text--------------------------
def get_list_table_index_by_alt_text(doc, alt_text):
    table_indices = []
    for i, table in enumerate(doc.tables):
        tblCaption = table._tbl.xpath('.//w:tblCaption')
        if tblCaption and tblCaption[0].get(qn('w:val')) == alt_text:
            table_indices.append(i)
    return table_indices
